# document_processor_agent.py

import docx
from typing import List, Dict, Any, Optional
from utils.azure_ai import AzureAIWrapper
import logging
import os
from io import BytesIO
import PyPDF2
from datetime import datetime
import json

logger = logging.getLogger(__name__)

class DocumentProcessorAgent:
    def __init__(self):
        self.ai_client = AzureAIWrapper()
        self.core_questions = []
        self.field_specific_questions = {}
        self.checklist_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'ethics_checklist.docx')
        
    def process_research_context(self, research_data: Dict[str, Any]) -> Dict[str, Any]:
        """Process research context and return relevant questions"""
        try:
            # Extract questions from ethics checklist
            questions = self.load_ethics_checklist()
            
            # Use AI to analyze research context and questions
            messages = [
                {
                    "role": "system",
                    "content": """You are an ethics review assistant. Analyze the research context and ethics questions 
                    to determine which questions are relevant and mandatory."""
                },
                {
                    "role": "user",
                    "content": f"""
                    Research Context:
                    {json.dumps(research_data, indent=2)}
                    
                    Ethics Questions:
                    {json.dumps(questions, indent=2)}
                    
                    For each question determine:
                    1. If it's a mandatory/core question (required for all research)
                    2. If it's field-specific and relevant to this research
                    3. If it requires document upload
                    4. If it should be included based on the research context
                    
                    Return the analysis in this JSON format:
                    {{
                        "mandatory": [
                            {{
                                "id": "question_id",
                                "question": "question_text",
                                "requires_document": true/false
                            }}
                        ],
                        "field_specific": [
                            {{
                                "id": "question_id",
                                "question": "question_text",
                                "requires_document": true/false,
                                "relevant": true/false
                            }}
                        ]
                    }}
                    """
                }
            ]
            
            analysis = self.ai_client.get_completion(messages)
            
            if not analysis or analysis == "AI service encountered an error.":
                return {
                    "status": "ERROR",
                    "message": "Unable to process research context due to AI service unavailability."
                }
            
            # Format questions based on AI analysis
            return self._format_questions(analysis)
            
        except Exception as e:
            logger.error(f"Error processing research context: {str(e)}")
            return {
                "status": "ERROR",
                "message": f"Error processing research context: {str(e)}"
            }

    def load_ethics_checklist(self) -> List[str]:
        """Load and parse the ethics checklist document"""
        try:
            doc = docx.Document(self.checklist_path)
            
            # Extract all questions from the document
            questions = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if self._is_question(text):
                    questions.append(text)
            
            return questions
            
        except Exception as e:
            logger.error(f"Error loading ethics checklist from {self.checklist_path}: {str(e)}")
            return []

    def _is_question(self, text: str) -> bool:
        """Check if the text is a question"""
        text = text.lower().strip()
        return (text.endswith('?') or 
                any(keyword in text for keyword in ['required', 'mandatory', 'must', 'should', 'please provide', 'please submit']))

    def _format_questions(self, ai_analysis: str) -> Dict[str, List[Dict]]:
        """Format questions based on AI analysis"""
        try:
            # Convert AI response to dictionary if it's a string
            if isinstance(ai_analysis, str):
                analysis = json.loads(ai_analysis)
            else:
                analysis = ai_analysis

            # Format mandatory questions
            formatted_mandatory = [
                {
                    'id': q.get('id', f'core_{i}'),
                    'question': q['question'],
                    'options': ['YES', 'NO'],
                    'is_core': True,
                    'requires_document': q.get('requires_document', False)
                } for i, q in enumerate(analysis.get('mandatory', []))
            ]

            # Format field-specific questions
            formatted_field_specific = [
                {
                    'id': q.get('id', f'field_{i}'),
                    'question': q['question'],
                    'options': ['YES', 'NO', 'N/A'],
                    'is_core': False,
                    'requires_document': q.get('requires_document', False)
                } for i, q in enumerate(analysis.get('field_specific', []))
                if q.get('relevant', True)  # Only include relevant questions
            ]

            return {
                'mandatory': formatted_mandatory,
                'field_specific': formatted_field_specific,
                'status': 'COMPLETED'
            }

        except Exception as e:
            logger.error(f"Error formatting questions: {str(e)}")
            return {
                'status': 'ERROR',
                'message': f"Error formatting questions: {str(e)}"
            }

    def review_document(self, content: bytes, filename: str, question: str) -> Dict[str, Any]:
        """Review a document for ethical considerations"""
        try:
            # Extract text from the document
            document_text = self._extract_document_text(content, filename)
            
            # Determine document type based on question and filename
            document_type = self._determine_document_type(question, filename)
            
            # Prepare messages for AI review with enhanced context
            messages = [
                {
                    "role": "system",
                    "content": """You are a document review assistant specializing in research ethics documentation.
                    Analyze the provided document in the context of the related ethics question.
                    Provide specific feedback on whether the document meets ethical requirements,
                    identify any missing elements, and suggest improvements."""
                },
                {
                    "role": "user",
                    "content": f"""
                    Ethics Question: {question}
                    
                    Document Name: {filename}
                    Document Type: {document_type}
                    
                    Document Content Preview:
                    {document_text[:3000] if len(document_text) > 3000 else document_text}
                    
                    Please analyze this document and provide:
                    1. An assessment of whether it adequately addresses the ethics question
                    2. Identification of any missing or incomplete elements
                    3. Specific recommendations for improvement
                    4. A status determination (APPROVED, ANALYZING, or NEEDS_REVISION)
                    
                    Format your response as a JSON object with these fields:
                    - status: "APPROVED", "ANALYZING", or "NEEDS_REVISION"
                    - analysis: A detailed analysis of the document
                    - missing_elements: List of any missing or incomplete elements
                    - recommendations: List of specific recommendations for improvement
                    - compliance_score: A number from 0-100 indicating how well the document meets requirements
                    """
                }
            ]
            
            # Call the AI service
            result = self.ai_client.get_completion(messages)
            
            if not result or result == "AI service encountered an error.":
                return {
                    "status": "ERROR",
                    "message": "Document review unavailable due to AI service error."
                }
            
            # Process the result
            if isinstance(result, str):
                try:
                    return json.loads(result)
                except json.JSONDecodeError:
                    return {
                        "status": "ERROR",
                        "message": "Unable to parse AI response for document review."
                    }
            else:
                return result
                
        except Exception as e:
            logger.error(f"Error reviewing document: {str(e)}")
            return {
                "status": "ERROR",
                "message": f"Error reviewing document: {str(e)}"
            }

    def _determine_document_type(self, question: str, filename: str) -> str:
        """Determine document type based on question and filename"""
        filename_lower = filename.lower()
        question_lower = question.lower()
        
        # Check for common document types in filename
        if "consent" in filename_lower:
            return "Consent Form"
        elif "protocol" in filename_lower:
            return "Research Protocol"
        elif "questionnaire" in filename_lower or "survey" in filename_lower:
            return "Survey/Questionnaire"
        elif "approval" in filename_lower or "letter" in filename_lower:
            return "Approval Letter"
        elif "cv" in filename_lower or "curriculum" in filename_lower or "resume" in filename_lower:
            return "CV/Resume"
        
        # Check for common document types in question
        if "consent" in question_lower:
            return "Consent Form"
        elif "protocol" in question_lower:
            return "Research Protocol"
        elif "questionnaire" in question_lower or "survey" in question_lower:
            return "Survey/Questionnaire"
        elif "approval" in question_lower or "letter" in question_lower:
            return "Approval Letter"
        elif "cv" in question_lower or "curriculum" in question_lower or "resume" in question_lower:
            return "CV/Resume"
        
        # Default document type
        return "Supporting Document"

    def _extract_document_text(self, content: bytes, filename: str) -> str:
        """Extract text from uploaded document"""
        try:
            if filename.lower().endswith('.pdf'):
                return self._extract_pdf_text(content)
            elif filename.lower().endswith(('.doc', '.docx')):
                return self._extract_word_text(content)
            else:
                raise ValueError(f"Unsupported file format: {filename}")
                
        except Exception as e:
            logger.error(f"Error extracting text from document: {str(e)}")
            return f"[Error extracting text from {filename}]"

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF document"""
        try:
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            return "[PDF text extraction error]"

    def _extract_word_text(self, content: bytes) -> str:
        """Extract text from Word document"""
        try:
            doc = docx.Document(BytesIO(content))
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error extracting Word text: {str(e)}")
            return "[Word document text extraction error]"

    def generate_review_report(self, reviews: Dict[str, Any]) -> Dict[str, Any]:
        """Generate a formatted review report using AI"""
        try:
            if not reviews:
                return {
                    "status": "ERROR",
                    "message": "No reviews available to generate report."
                }
            
            # Prepare messages for AI to generate report
            messages = [
                {
                    "role": "system",
                    "content": """You are an ethics review report generator. Create a comprehensive, 
                    well-formatted report based on the document reviews provided."""
                },
                {
                    "role": "user",
                    "content": f"""
                    Document Reviews:
                    {json.dumps(reviews, indent=2)}
                    
                    Generate a comprehensive ethics review report that includes:
                    1. An executive summary of the overall review status
                    2. Detailed analysis of each document
                    3. Identification of any missing elements or issues
                    4. Specific recommendations for improvement
                    5. Next steps for the researcher
                    
                    Format the report in a professional, well-structured manner.
                    """
                }
            ]
            
            # Call the AI service
            report_content = self.ai_client.get_completion(messages)
            
            if not report_content or report_content == "AI service encountered an error.":
                return {
                    "status": "ERROR",
                    "message": "Report generation unavailable due to AI service error."
                }
            
            return {
                "status": "COMPLETED",
                "report": report_content,
                "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
        except Exception as e:
            logger.error(f"Error generating review report: {str(e)}")
            return {
                "status": "ERROR",
                "message": f"Error generating review report: {str(e)}"
            }
        
    def get_document_requirements(self, question: Dict[str, Any]) -> Dict[str, Any]:
        """Get document requirements for a specific question"""
        try:
            # Determine document type based on question content
            question_text = question['question']
            
            # Prepare messages for AI to determine document requirements
            messages = [
                {
                    "role": "system",
                    "content": """You are an ethics review assistant specializing in research documentation requirements.
                    Determine what documents are required for the given ethics question and what elements those documents should contain."""
                },
                {
                    "role": "user",
                    "content": f"""
                    Ethics Question: {question_text}
                    
                    Please determine:
                    1. If this question requires supporting documentation
                    2. What type of document is required
                    3. What specific elements should be included in the document
                    4. Any formatting or signature requirements
                    
                    Format your response as a JSON object with these fields:
                    - requires_document: true/false
                    - document_type: The type of document required
                    - required_elements: List of specific elements that must be included
                    - format_requirements: List of formatting or signature requirements
                    """
                }
            ]
            
            # Call the AI service
            result = self.ai_client.get_completion(messages)
            
            if not result or result == "AI service encountered an error.":
                return {
                    "status": "ERROR",
                    "message": "Document requirements unavailable due to AI service error."
                }
            
            # Process the result
            if isinstance(result, str):
                try:
                    parsed_result = json.loads(result)
                    parsed_result["status"] = "COMPLETED"
                    return parsed_result
                except json.JSONDecodeError:
                    return {
                        "status": "ERROR",
                        "message": "Unable to parse AI response for document requirements."
                    }
            else:
                result["status"] = "COMPLETED"
                return result
                
        except Exception as e:
            logger.error(f"Error getting document requirements: {str(e)}")
            return {
                "status": "ERROR",
                "message": f"Error getting document requirements: {str(e)}"
            }
    def validate_document(self, document_content: bytes, document_name: str, requirements: Dict[str, Any]) -> Dict[str, Any]:
        """Validate if uploaded document meets requirements"""
        try:
            # Extract text from the document
            document_text = self._extract_document_text(document_content, document_name)
            
            # Get document requirements
            document_type = requirements.get('document_type', 'Supporting Document')
            required_elements = requirements.get('required_elements', [])
            format_requirements = requirements.get('format_requirements', [])
            
            # Prepare messages for AI validation
            messages = [
                {
                    "role": "system",
                    "content": """You are a document validation assistant specializing in research ethics documentation.
                    Validate if the provided document meets the specified requirements."""
                },
                {
                    "role": "user",
                    "content": f"""
                    Document Name: {document_name}
                    Document Type: {document_type}
                    
                    Document Content Preview:
                    {document_text[:3000] if len(document_text) > 3000 else document_text}
                    
                    Required Elements:
                    {json.dumps(required_elements)}
                    
                    Format Requirements:
                    {json.dumps(format_requirements)}
                    
                    Please validate this document and provide:
                    1. Whether the document is valid for its intended purpose
                    2. Which required elements are present
                    3. Which required elements are missing
                    4. Specific feedback for improvement
                    5. A status determination (APPROVED, ANALYZING, or NEEDS_REVISION)
                    
                    Format your response as a JSON object with these fields:
                    - valid: true/false
                    - present_elements: List of required elements that are present
                    - missing_elements: List of required elements that are missing
                    - format_issues: List of any format requirement issues
                    - feedback: Detailed feedback for improvement
                    - status: "APPROVED", "ANALYZING", or "NEEDS_REVISION"
                    """
                }
            ]
            
            # Call the AI service
            result = self.ai_client.get_completion(messages)
            
            if not result or result == "AI service encountered an error.":
                return {
                    "status": "ERROR",
                    "message": "Document validation unavailable due to AI service error."
                }
            
            # Process the result
            if isinstance(result, str):
                try:
                    parsed_result = json.loads(result)
                    return parsed_result
                except json.JSONDecodeError:
                    return {
                        "status": "ERROR",
                        "message": "Unable to parse AI response for document validation."
                    }
            else:
                return result
                
        except Exception as e:
            logger.error(f"Error validating document: {str(e)}")
            return {
                "status": "ERROR",
                "message": f"Error validating document: {str(e)}"
            }
    
    def generate_ethics_approval_letter(self, application_data: Dict[str, Any], 
                                        reviews: Dict[str, Any]) -> Dict[str, Any]:
        """Generate an ethics approval letter based on application and reviews"""
        try:
            # Extract researcher information
            researcher_name = application_data.get('researcher_name', 'Researcher')
            project_title = application_data.get('project_title', 'Research Project')
            
            # Prepare messages for AI to generate letter
            messages = [
                {
                    "role": "system",
                    "content": """You are an ethics committee assistant responsible for drafting 
                    approval letters. Create a formal ethics review letter based on the provided 
                    application data and review outcomes."""
                },
                {
                    "role": "user",
                    "content": f"""
                    Application Data:
                    {json.dumps(application_data, indent=2)}
                    
                    Review Outcomes:
                    {json.dumps(reviews, indent=2)}
                    
                    Generate a formal ethics committee letter addressing the researcher.
                    The letter should include:
                    1. Formal greeting and introduction
                    2. Reference to the submitted application
                    3. The committee's decision based on the reviews
                    4. Summary of key findings from the reviews
                    5. Next steps or requirements (if any)
                    6. Closing with contact information for questions
                    
                    Format the letter professionally with appropriate sections and paragraphs.
                    """
                }
            ]
            
            # Generate letter content
            letter_content = self.ai_client.get_completion(messages)
            
            if not letter_content or letter_content == "AI service encountered an error.":
                return {
                    'status': "ERROR",
                    'message': "Unable to generate approval letter due to AI service error."
                }
            
            # Create letter document
            letter_doc = docx.Document()
            
            # Add header
            header = letter_doc.add_heading('Ethics Review Committee', 0)
            header.alignment = 1  # Center alignment
            
            # Add date
            date_paragraph = letter_doc.add_paragraph()
            date_paragraph.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
            
            # Add letter content
            for paragraph in letter_content.split('\n\n'):
                if paragraph.strip():
                    letter_doc.add_paragraph(paragraph.strip())
            
            # Save to BytesIO
            letter_bytes = BytesIO()
            letter_doc.save(letter_bytes)
            letter_bytes.seek(0)
            
            return {
                'content': letter_bytes.getvalue(),
                'filename': f"ethics_review_{datetime.now().strftime('%Y%m%d')}.docx",
                'status': 'COMPLETED'
            }
            
        except Exception as e:
            logger.error(f"Error generating approval letter: {str(e)}")
            return {
                'status': "ERROR",
                'message': f"Error generating approval letter: {str(e)}"
            }

    def analyze_research_risks(self, research_data: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze potential ethical risks in the research proposal"""
        try:
            # Prepare messages for AI risk analysis
            messages = [
                {
                    "role": "system",
                    "content": """You are an ethics risk assessment specialist. Analyze the provided research 
                    data to identify potential ethical risks and provide recommendations."""
                },
                {
                    "role": "user",
                    "content": f"""
                    Research Data:
                    {json.dumps(research_data, indent=2)}
                    
                    Please analyze this research proposal for potential ethical risks in these categories:
                    1. Participant safety and wellbeing
                    2. Data privacy and security
                    3. Informed consent
                    4. Vulnerable populations
                    5. Conflicts of interest
                    6. Research integrity
                    
                    For each identified risk:
                    - Provide a risk level (Low, Medium, High)
                    - Explain the potential impact
                    - Suggest mitigation strategies
                    
                    Return your analysis as a structured JSON object.
                    """
                }
            ]
            
            # Get risk analysis from AI
            risk_analysis = self.ai_client.get_completion(messages)
            
            if not risk_analysis or risk_analysis == "AI service encountered an error.":
                return {
                    "status": "ERROR",
                    "message": "Risk analysis unavailable due to AI service error."
                }
            
            # Parse the response
            if isinstance(risk_analysis, str):
                try:
                    parsed_result = json.loads(risk_analysis)
                    parsed_result["status"] = "COMPLETED"
                    return parsed_result
                except json.JSONDecodeError:
                    return {
                        "status": "COMPLETED",
                        "analysis": risk_analysis
                    }
            else:
                risk_analysis["status"] = "COMPLETED"
                return risk_analysis
                
        except Exception as e:
            logger.error(f"Error analyzing research risks: {str(e)}")
            return {
                "status": "ERROR",
                "message": f"Error analyzing research risks: {str(e)}"
            }

    def generate_consent_form_template(self, research_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate a consent form template based on research data"""
        try:
            # Extract relevant information
            project_title = research_data.get('project_title', 'Research Project')
            researcher_name = research_data.get('researcher_name', 'Researcher')
            institution = research_data.get('institution', 'Institution')
            
            # Prepare messages for AI to generate consent form
            messages = [
                {
                    "role": "system",
                    "content": """You are a research ethics specialist who creates consent form templates. 
                    Generate a comprehensive consent form based on the provided research details."""
                },
                {
                    "role": "user",
                    "content": f"""
                    Research Details:
                    {json.dumps(research_data, indent=2)}
                    
                    Create a comprehensive consent form template that includes:
                    1. Project title and researcher information
                    2. Purpose of the research
                    3. What participation involves
                    4. Risks and benefits
                    5. Confidentiality and data protection
                    6. Voluntary participation and right to withdraw
                    7. Contact information
                    8. Consent declaration and signature blocks
                    
                    Format the content as a professional document with clear sections.
                    """
                }
            ]
            
            # Generate consent form content
            consent_content = self.ai_client.get_completion(messages)
            
            if not consent_content or consent_content == "AI service encountered an error.":
                return {
                    'status': "ERROR",
                    'message': "Unable to generate consent form template due to AI service error."
                }
            
            # Create consent form document
            consent_doc = docx.Document()
            
            # Add title
            title = consent_doc.add_heading('INFORMED CONSENT FORM', 0)
            title.alignment = 1  # Center alignment
            
            # Add project title
            project_heading = consent_doc.add_heading(f"Project: {project_title}", 1)
            
            # Add consent content
            for paragraph in consent_content.split('\n\n'):
                if paragraph.strip():
                    consent_doc.add_paragraph(paragraph.strip())
            
            # Save to BytesIO
            consent_bytes = BytesIO()
            consent_doc.save(consent_bytes)
            consent_bytes.seek(0)
            
            return {
                'content': consent_bytes.getvalue(),
                'filename': f"consent_form_template_{datetime.now().strftime('%Y%m%d')}.docx",
                'status': 'COMPLETED'
            }
            
        except Exception as e:
            logger.error(f"Error generating consent form template: {str(e)}")
            return {
                'status': "ERROR",
                'message': f"Error generating consent form template: {str(e)}"
            }

    def check_application_completeness(self, application_data: Dict[str, Any]) -> Dict[str, Any]:
        """Check if an ethics application is complete and ready for review"""
        try:
            # Define required fields for a complete application
            required_fields = [
                'researcher_name', 
                'project_title', 
                'institution',
                'research_summary',
                'start_date',
                'end_date'
            ]
            
            # Prepare messages for AI to check completeness
            messages = [
                {
                    "role": "system",
                    "content": """You are an ethics application validator. Analyze the provided application 
                    data to determine if it's complete and ready for review."""
                },
                {
                    "role": "user",
                    "content": f"""
                    Application Data:
                    {json.dumps(application_data, indent=2)}
                    
                    Required Fields:
                    {json.dumps(required_fields)}
                    
                    Please analyze this application and determine:
                    1. If all required fields are completed
                    2. If all required documents are attached
                    3. If all questions have been answered
                    4. The overall completeness percentage
                    5. Any specific issues that need to be addressed
                    
                    Return your analysis as a structured JSON object with these fields:
                    - is_complete: true/false
                    - completion_percentage: number from 0-100
                    - missing_fields: list of missing required fields
                    - missing_documents: list of missing required documents
                    - unanswered_questions: list of unanswered questions
                    - recommendations: list of specific recommendations to complete the application
                    """
                }
            ]
            
            # Get completeness check from AI
            completeness_check = self.ai_client.get_completion(messages)
            
            if not completeness_check or completeness_check == "AI service encountered an error.":
                return {
                    "status": "ERROR",
                    "message": "Application completeness check unavailable due to AI service error."
                }
            
            # Parse the response
            if isinstance(completeness_check, str):
                try:
                    parsed_result = json.loads(completeness_check)
                    parsed_result["status"] = "COMPLETED"
                    return parsed_result
                except json.JSONDecodeError:
                    return {
                        "status": "ERROR",
                        "message": "Unable to parse AI response for application completeness check."
                    }
            else:
                completeness_check["status"] = "COMPLETED"
                return completeness_check
                
        except Exception as e:
            logger.error(f"Error checking application completeness: {str(e)}")
            return {
                "status": "ERROR",
                "message": f"Error checking application completeness: {str(e)}"
            }  